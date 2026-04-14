"""Export service — generates transcription exports in various formats.

Supported formats:
  txt   — plain text with speaker labels and timestamps
  srt   — SubRip subtitle format
  vtt   — WebVTT subtitle format
  md    — Markdown document
  docx  — Microsoft Word document
"""

from __future__ import annotations

import io
import math
from typing import Any, Optional


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _resolve_speaker(raw: str, speakers: dict[str, str]) -> str:
    """Return human-readable speaker name, falling back to the raw label."""
    return speakers.get(raw, raw) if speakers else raw


def _seconds_to_hms(seconds: float) -> str:
    """Convert seconds to HH:MM:SS."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _seconds_to_srt_ts(seconds: float) -> str:
    """Convert seconds to SRT timestamp HH:MM:SS,mmm."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = math.floor(seconds % 60)
    ms = round((seconds - math.floor(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _seconds_to_vtt_ts(seconds: float) -> str:
    """Convert seconds to VTT timestamp HH:MM:SS.mmm."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = math.floor(seconds % 60)
    ms = round((seconds - math.floor(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def _title(transcription: Any) -> str:
    return transcription.title or transcription.original_filename


# ─────────────────────────────────────────────────────────────────────────────
# TXT
# ─────────────────────────────────────────────────────────────────────────────

def export_txt(transcription: Any) -> bytes:
    """Plain text export — speaker labels + timestamps per segment, or raw text for whisper."""
    speakers: dict = transcription.speakers or {}
    segments: list = transcription.segments or []
    lines: list[str] = []

    lines.append(_title(transcription))
    lines.append("=" * len(_title(transcription)))
    lines.append("")

    if segments:
        prev_speaker = None
        for seg in segments:
            speaker_raw = seg.get("speaker")
            speaker = _resolve_speaker(speaker_raw, speakers) if speaker_raw else None
            start = _seconds_to_hms(seg.get("start", 0))
            text = seg.get("text", "").strip()
            if not text:
                continue
            if speaker and speaker != prev_speaker:
                lines.append(f"[{start}] {speaker}")
                prev_speaker = speaker
            elif not speaker:
                lines.append(f"[{start}]")
            lines.append(text)
            lines.append("")
    elif transcription.text:
        lines.append(transcription.text)

    return "\n".join(lines).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# SRT
# ─────────────────────────────────────────────────────────────────────────────

def export_srt(transcription: Any) -> bytes:
    """SubRip (.srt) subtitle export."""
    speakers: dict = transcription.speakers or {}
    segments: list = transcription.segments or []
    lines: list[str] = []

    for i, seg in enumerate(segments, start=1):
        start = _seconds_to_srt_ts(seg.get("start", 0))
        end = _seconds_to_srt_ts(seg.get("end", seg.get("start", 0) + 1))
        text = seg.get("text", "").strip()
        if not text:
            continue
        speaker_raw = seg.get("speaker")
        if speaker_raw:
            speaker = _resolve_speaker(speaker_raw, speakers)
            text = f"[{speaker}] {text}"
        lines.append(str(i))
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# VTT
# ─────────────────────────────────────────────────────────────────────────────

def export_vtt(transcription: Any) -> bytes:
    """WebVTT (.vtt) subtitle export."""
    speakers: dict = transcription.speakers or {}
    segments: list = transcription.segments or []
    lines: list[str] = ["WEBVTT", ""]

    for seg in segments:
        start = _seconds_to_vtt_ts(seg.get("start", 0))
        end = _seconds_to_vtt_ts(seg.get("end", seg.get("start", 0) + 1))
        text = seg.get("text", "").strip()
        if not text:
            continue
        speaker_raw = seg.get("speaker")
        if speaker_raw:
            speaker = _resolve_speaker(speaker_raw, speakers)
            text = f"<v {speaker}>{text}"
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# Markdown
# ─────────────────────────────────────────────────────────────────────────────

def export_md(transcription: Any, summaries: Optional[list] = None) -> bytes:
    """Markdown export — structured document with metadata, transcript, and summaries."""
    speakers: dict = transcription.speakers or {}
    segments: list = transcription.segments or []
    is_whisper = str(transcription.recording_type) in ("whisper", "RecordingType.whisper")
    lines: list[str] = []

    # Header
    lines.append(f"# {_title(transcription)}")
    lines.append("")

    # Metadata
    rec_type = "Whisper memo" if is_whisper else "Recording"
    lines.append(f"**Type:** {rec_type}  ")
    if transcription.audio_duration:
        lines.append(f"**Duration:** {_seconds_to_hms(transcription.audio_duration)}  ")
    if transcription.language:
        lines.append(f"**Language:** {transcription.language}  ")
    lines.append("")

    # Summaries (if any)
    if summaries:
        for summary in summaries:
            lines.append("---")
            lines.append("")
            lines.append("## Summary")
            lines.append("")
            lines.append(summary.content)
            lines.append("")

    lines.append("---")
    lines.append("")

    # Transcript body
    if is_whisper:
        lines.append("## Transcript")
        lines.append("")
        lines.append(transcription.text or "")
    else:
        lines.append("## Transcript")
        lines.append("")
        if segments:
            prev_speaker = None
            for seg in segments:
                speaker_raw = seg.get("speaker")
                speaker = _resolve_speaker(speaker_raw, speakers) if speaker_raw else None
                start = _seconds_to_hms(seg.get("start", 0))
                text = seg.get("text", "").strip()
                if not text:
                    continue
                if speaker and speaker != prev_speaker:
                    lines.append(f"**[{start}] {speaker}**")
                    prev_speaker = speaker
                lines.append(text)
                lines.append("")
        elif transcription.text:
            lines.append(transcription.text)

    return "\n".join(lines).encode("utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def export_docx(transcription: Any, summaries: Optional[list] = None) -> bytes:
    """Microsoft Word (.docx) export."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──
    title_para = doc.add_heading(_title(transcription), level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Metadata table ──
    is_whisper = str(transcription.recording_type) in ("whisper", "RecordingType.whisper")
    rec_type = "Whisper memo" if is_whisper else "Recording"
    meta_lines = [f"Type: {rec_type}"]
    if transcription.audio_duration:
        meta_lines.append(f"Duration: {_seconds_to_hms(transcription.audio_duration)}")
    if transcription.language:
        meta_lines.append(f"Language: {transcription.language}")

    meta_para = doc.add_paragraph()
    for line in meta_lines:
        run = meta_para.add_run(line + "\n")
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Summaries ──
    if summaries:
        for summary in summaries:
            doc.add_heading("Summary", level=2)
            _add_markdown_paragraphs(doc, summary.content)
            doc.add_paragraph()

    # ── Transcript ──
    doc.add_heading("Transcript", level=2)

    speakers: dict = transcription.speakers or {}
    segments: list = transcription.segments or []

    if is_whisper:
        doc.add_paragraph(transcription.text or "")
    elif segments:
        prev_speaker = None
        for seg in segments:
            speaker_raw = seg.get("speaker")
            speaker = _resolve_speaker(speaker_raw, speakers) if speaker_raw else None
            start = _seconds_to_hms(seg.get("start", 0))
            text = seg.get("text", "").strip()
            if not text:
                continue
            if speaker and speaker != prev_speaker:
                speaker_para = doc.add_paragraph()
                run = speaker_para.add_run(f"[{start}]  {speaker}")
                run.bold = True
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
                run.font.size = Pt(10)
                prev_speaker = speaker
            doc.add_paragraph(text)
    elif transcription.text:
        doc.add_paragraph(transcription.text)

    # ── Save to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_markdown_paragraphs(doc: Any, content: str) -> None:
    """Naive markdown-to-docx renderer for summary content."""
    from docx.shared import Pt, RGBColor

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=4)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=3)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped[3] == "x"
            text = stripped[6:]
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(("☑ " if checked else "☐ ") + text)
            run.font.size = Pt(11)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:])
        elif stripped.startswith("| "):
            # Simple table row — just render as plain text
            doc.add_paragraph(stripped)
        else:
            # Handle inline bold (**text**)
            p = doc.add_paragraph()
            _render_inline(p, stripped)


def _render_inline(para: Any, text: str) -> None:
    """Render inline markdown bold/italic into a paragraph."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)
